# build_complete.py - исправленная сборка с включением модуля email
import os
import shutil
import subprocess
import sys
from pathlib import Path
import json


def install_pyinstaller():
    """Установить PyInstaller если не установлен"""
    try:
        import PyInstaller
        print("✅ PyInstaller уже установлен")
        return True
    except ImportError:
        print("Установка PyInstaller...")
        try:
            subprocess.run([sys.executable, '-m', 'pip', 'install', 'pyinstaller'],
                           check=True, capture_output=True, text=True)
            print("✅ PyInstaller успешно установлен")
            return True
        except subprocess.CalledProcessError as e:
            print(f"❌ Ошибка установки PyInstaller: {e}")
            if e.stderr:
                print(f"Детали: {e.stderr}")
            return False


def build_complete():
    print("🚀 Запуск исправленной сборки с включением модуля email...")
    print(f"Текущая директория: {os.getcwd()}")

    # Проверяем существование основных файлов
    required_files = ['main.py', 'main_window.py', 'settings.py', 'theme_manager.py',
                      'license_manager.py', 'update_manager.py', 'widgets.py']

    missing_files = []
    for file in required_files:
        if not os.path.exists(file):
            missing_files.append(file)

    if missing_files:
        print(f"❌ Отсутствуют файлы: {missing_files}")
        return

    # Установка PyInstaller если нужно
    if not install_pyinstaller():
        print("Создаем портативное решение...")
        create_portable_solution()
        return

    # Создаем недостающие конфиги
    create_missing_configs()

    # Очистка предыдущих сборок
    for folder in ['build', 'dist']:
        if os.path.exists(folder):
            shutil.rmtree(folder)
            print(f"🗑️ Очищена папка: {folder}")

    # Создаем правильный spec файл с включением email
    create_correct_spec_file()

    # Запускаем сборку
    try:
        print("🔨 Запуск PyInstaller...")
        result = subprocess.run([
            sys.executable, '-m', 'PyInstaller',
            'document_filler.spec', '--clean', '--noconfirm'
        ], check=True, capture_output=True, text=True)

        print("✅ Сборка завершена успешно!")

        # Проверяем результат сборки
        if check_build_result():
            create_final_distribution()
        else:
            print("❌ Сборка не создала исполняемый файл")
            create_backup_solution()

    except subprocess.CalledProcessError as e:
        print(f"❌ Ошибка сборки: {e}")
        if e.stderr:
            print(f"Детали ошибки: {e.stderr}")
        create_backup_solution()


def create_missing_configs():
    """Создать недостающие конфигурационные файлы"""
    configs = {
        'version_config.json': {
            "current_version": "1.0.0",
            "update_url": "",
            "check_updates_on_start": False,
            "update_channel": "stable"
        },
        'repo_config.json': {
            "type": "github",
            "owner": "your-username",
            "repo": "your-repo-name",
            "branch": "main",
            "token": ""
        }
    }

    for filename, config in configs.items():
        if not os.path.exists(filename):
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
            print(f"📄 Создан {filename}")


def create_correct_spec_file():
    """Создать правильный spec файл с включением модуля email"""

    # Проверяем существование папки шаблонов
    templates_exists = os.path.exists('Шаблоны')
    datas = []

    # Обязательные файлы
    datas.append("('version_config.json', '.')")
    datas.append("('repo_config.json', '.')")

    # Папка шаблонов
    if templates_exists:
        datas.append("('Шаблоны', 'Шаблоны')")
        print("✅ Папка 'Шаблоны' будет включена в сборку")
    else:
        print("⚠️ Папка 'Шаблоны' не найдена, создаем пустую")
        os.makedirs('Шаблоны', exist_ok=True)
        datas.append("('Шаблоны', 'Шаблоны')")

    spec_content = f'''# -*- mode: python ; coding: utf-8 -*-

block_cipher = None

a = Analysis(
    ['main.py'],
    pathex=[],
    binaries=[],
    datas=[
        {', '.join(datas)}
    ],
    hiddenimports=[
        # Основные модули приложения
        'main_window', 'settings', 'theme_manager', 
        'license_manager', 'update_manager', 'widgets',

        # PyQt5 модули
        'PyQt5', 'PyQt5.QtCore', 'PyQt5.QtGui', 'PyQt5.QtWidgets', 'PyQt5.QtNetwork',
        'PyQt5.sip',

        # Документы и шаблоны
        'openpyxl', 'docxtpl', 'jinja2', 'docx',
        'lxml', 'lxml.etree', 'lxml._elementpath',

        # Сеть - ВАЖНО: включаем все необходимые модули для requests
        'requests', 'urllib3', 'chardet', 'idna', 'certifi',
        'email', 'email.mime', 'email.mime.text', 'email.mime.multipart',
        'email.mime.base', 'email.encoders', 'email.utils',
        'ssl', 'http', 'http.client', 'http.cookies',

        # Системные модули которые нужны для работы
        'hashlib', 'json', 'datetime', 'os', 'sys', 're',
        'uuid', 'platform', 'threading', 'tempfile', 'zipfile',
        'xml', 'xml.etree', 'xml.etree.ElementTree'
    ],
    hookspath=[],
    hooksconfig={{}},
    runtime_hooks=[],
    excludes=[
        # Исключаем только действительно ненужные модули
        'tkinter', 'unittest', 'test', 'pydoc',
        'numpy', 'pandas', 'scipy', 'matplotlib', 'PIL',
        'pygame', 'wx', 'gtk', 'curses',
        'concurrent', 'distutils', 'setuptools',
        'pip', 'wheel', 'pkg_resources', 'notebook',
        'jupyter', 'ipython', 'qtpy', 'pyqtgraph'
    ],
    win_no_prefer_redirects=False,
    win_private_assemblies=False,
    cipher=block_cipher,
    noarchive=False,
    optimize=1,
)

pyz = PYZ(a.pure, a.zipped_data, cipher=block_cipher)

exe = EXE(
    pyz,
    a.scripts,
    a.binaries,
    a.zipfiles,
    a.datas,
    [],
    name='DocumentFiller',
    debug=False,
    bootloader_ignore_signals=False,
    strip=False,
    upx=True,
    upx_exclude=[],
    runtime_tmpdir=None,
    console=False,
    disable_windowed_traceback=False,
    argv_emulation=False,
    target_arch=None,
    codesign_identity=None,
    entitlements_file=None,
    icon=None,
)

coll = COLLECT(
    exe,
    a.binaries,
    a.zipfiles,
    a.datas,
    strip=False,
    upx=True,
    upx_exclude=[],
    name='DocumentFiller'
)
'''

    with open('document_filler.spec', 'w', encoding='utf-8') as f:
        f.write(spec_content)
    print("✅ Создан исправленный spec файл с включением модуля email")


def check_build_result():
    """Проверить результат сборки"""
    dist_dir = 'dist/DocumentFiller'

    if not os.path.exists(dist_dir):
        print(f"❌ Папка сборки не найдена: {dist_dir}")
        return False

    exe_path = os.path.join(dist_dir, 'DocumentFiller.exe')
    if not os.path.exists(exe_path):
        print(f"❌ Исполняемый файл не найден: {exe_path}")
        return False

    print(f"✅ Исполняемый файл создан: {exe_path}")

    # Проверяем размер EXE файла
    exe_size = os.path.getsize(exe_path) / (1024 * 1024)
    print(f"📦 Размер EXE файла: {exe_size:.2f} MB")

    return True


def create_final_distribution():
    """Создать финальную дистрибуцию"""
    dist_dir = 'DocumentFiller_Final'

    # Очистка предыдущей сборки
    if os.path.exists(dist_dir):
        shutil.rmtree(dist_dir)
        print(f"🗑️ Очищена предыдущая сборка: {dist_dir}")

    # Копируем собранное приложение
    source_dir = 'dist/DocumentFiller'
    if os.path.exists(source_dir):
        print(f"📁 Копируем сборку из: {source_dir}")
        shutil.copytree(source_dir, dist_dir)

        # Проверяем содержимое финальной сборки
        check_final_distribution(dist_dir)

        print(f"🎉 ФИНАЛЬНАЯ СБОРКА СОЗДАНА: {dist_dir}")

        # Показываем размер сборки
        total_size = sum(f.stat().st_size for f in Path(dist_dir).rglob('*') if f.is_file())
        print(f"📦 Общий размер сборки: {total_size / (1024 * 1024):.2f} MB")

        # Показываем содержимое
        print("📋 Содержимое сборки:")
        for root, dirs, files in os.walk(dist_dir):
            level = root.replace(dist_dir, '').count(os.sep)
            indent = ' ' * 2 * level
            print(f'{indent}{os.path.basename(root)}/')
            subindent = ' ' * 2 * (level + 1)
            for file in files:
                file_size = os.path.getsize(os.path.join(root, file)) / 1024
                print(f'{subindent}{file} ({file_size:.1f} KB)')
    else:
        print(f"❌ Папка сборки не найдена: {source_dir}")


def check_final_distribution(dist_dir):
    """Проверить финальную дистрибуцию"""
    print("🔍 Проверка финальной сборки...")

    # Проверяем основные файлы
    required_files = [
        'DocumentFiller.exe',
        'version_config.json',
        'repo_config.json'
    ]

    for file in required_files:
        file_path = os.path.join(dist_dir, file)
        if os.path.exists(file_path):
            print(f"✅ Найден: {file}")
        else:
            print(f"❌ Отсутствует: {file}")

    # Проверяем папку шаблонов
    templates_dir = os.path.join(dist_dir, 'Шаблоны')
    if os.path.exists(templates_dir):
        template_files = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
        print(f"✅ Папка 'Шаблоны' содержит {len(template_files)} .docx файлов")

        if not template_files:
            print("⚠️ В папке 'Шаблоны' нет .docx файлов!")
            # Создаем тестовый шаблон
            create_sample_template(templates_dir)
    else:
        print("❌ Папка 'Шаблоны' не найдена в сборке!")
        # Создаем пустую папку шаблонов
        os.makedirs(templates_dir, exist_ok=True)
        create_sample_template(templates_dir)

    # Тестируем запуск приложения
    test_application_launch(dist_dir)

    create_readme(dist_dir)


def test_application_launch(dist_dir):
    """Протестировать запуск приложения"""
    print("🧪 Тестирование запуска приложения...")

    exe_path = os.path.join(dist_dir, 'DocumentFiller.exe')

    if not os.path.exists(exe_path):
        print("❌ Исполняемый файл не найден для тестирования")
        return

    try:
        # Запускаем приложение на 3 секунды чтобы проверить ошибки импорта
        import subprocess
        import time

        print("🔄 Запускаем приложение для проверки импортов...")
        process = subprocess.Popen([exe_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)

        # Ждем 5 секунд
        time.sleep(5)

        # Завершаем процесс
        process.terminate()
        try:
            process.wait(timeout=5)
        except subprocess.TimeoutExpired:
            process.kill()

        stdout, stderr = process.communicate()

        if stderr:
            error_output = stderr.decode('utf-8', errors='ignore')
            if "ModuleNotFoundError" in error_output or "ImportError" in error_output:
                print("❌ Обнаружены ошибки импорта:")
                print(error_output)
            else:
                print("✅ Приложение запустилось без критических ошибок импорта")
        else:
            print("✅ Приложение успешно запустилось!")

    except Exception as e:
        print(f"⚠️ Ошибка при тестировании запуска: {e}")


def create_sample_template(templates_dir):
    """Создать пример шаблона если папка пустая"""
    try:
        # Пытаемся использовать python-docx если установлен
        from docx import Document

        doc = Document()
        doc.add_heading('Шаблон документа DocumentFiller', 0)
        doc.add_paragraph('Это тестовый шаблон документа для программы DocumentFiller.')
        doc.add_paragraph('Доступные поля для заполнения:')

        fields = [
            ('Фамилия', '{{ n }}'),
            ('Имя', '{{ fn }}'),
            ('Отчество', '{{ mn }}'),
            ('Регистрация', '{{ reg }}'),
            ('Серия паспорта', '{{ ps }}'),
            ('Номер паспорта', '{{ pn }}'),
            ('Паспорт выдан', '{{ pi }}'),
            ('Дата выдачи', '{{ di }}'),
            ('Серия УЧО', '{{ cs }}'),
            ('Номер УЧО', '{{ cn }}'),
            ('Текущая дата', '{{ current_date }}')
        ]

        for field_name, field_code in fields:
            doc.add_paragraph(f'{field_name}: {field_code}')

        template_path = os.path.join(templates_dir, 'пример_шаблона.docx')
        doc.save(template_path)
        print(f"📄 Создан пример шаблона: {template_path}")
    except Exception as e:
        print(f"⚠️ Не удалось создать пример шаблона: {e}")
        # Создаем простой текстовый файл как запасной вариант
        template_content = '''Шаблон документа DocumentFiller

Доступные поля для заполнения:
Фамилия: {{ n }}
Имя: {{ fn }} 
Отчество: {{ mn }}
Регистрация: {{ reg }}
Серия паспорта: {{ ps }}
Номер паспорта: {{ pn }}
Паспорт выдан: {{ pi }}
Дата выдачи: {{ di }}
Серия УЧО: {{ cs }}
Номер УЧО: {{ cn }}
Текущая дата: {{ current_date }}

Сохраните этот файл как .docx для использования в программе.'''

        template_path = os.path.join(templates_dir, 'пример_шаблона.txt')
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(template_content)
        print(f"📄 Создан текстовый пример шаблона: {template_path}")


def create_backup_solution():
    """Создать резервное решение через командную строку"""
    print("🛡️ Создаем резервное решение через командную строку...")

    try:
        # Собираем команду с включением всех необходимых модулей
        cmd = [
            sys.executable, '-m', 'PyInstaller',
            'main.py',
            '--name=DocumentFiller',
            '--windowed',
            '--onedir',
            '--clean',
            '--noconfirm',
            '--add-data=version_config.json:.',
            '--add-data=repo_config.json:.',
            '--hidden-import=email',
            '--hidden-import=email.mime',
            '--hidden-import=email.mime.text',
            '--hidden-import=email.mime.multipart',
            '--hidden-import=email.mime.base',
            '--hidden-import=email.encoders',
            '--hidden-import=email.utils',
        ]

        # Добавляем папку шаблонов
        if os.path.exists('Шаблоны'):
            cmd.append('--add-data=Шаблоны:Шаблоны')

        print(f"🔧 Команда: {' '.join(cmd)}")
        result = subprocess.run(cmd, check=True, capture_output=True, text=True)

        print("✅ Резервная сборка завершена!")

        if check_build_result():
            create_final_distribution()
        else:
            create_portable_solution()

    except Exception as e:
        print(f"❌ Резервная сборка не удалась: {e}")
        create_portable_solution()


def create_portable_solution():
    """Создать портативное решение"""
    print("💼 Создаем портативное решение...")

    portable_dir = 'DocumentFiller_Portable'

    # Очистка
    if os.path.exists(portable_dir):
        shutil.rmtree(portable_dir)

    os.makedirs(portable_dir)
    print(f"📁 Создана папка: {portable_dir}")

    # Копируем ВСЕ Python файлы
    python_files = [f for f in os.listdir('.') if f.endswith('.py')]
    copied_count = 0

    for file in python_files:
        try:
            shutil.copy2(file, portable_dir)
            copied_count += 1
        except Exception as e:
            print(f"⚠️ Не удалось скопировать {file}: {e}")

    print(f"📄 Скопировано {copied_count} Python файлов")

    # Копируем конфиги
    config_files = ['version_config.json', 'repo_config.json']
    for config in config_files:
        if os.path.exists(config):
            shutil.copy2(config, portable_dir)
            print(f"📄 Скопирован: {config}")

    # Копируем папку шаблонов
    if os.path.exists('Шаблоны'):
        templates_dest = os.path.join(portable_dir, 'Шаблоны')
        shutil.copytree('Шаблоны', templates_dest)
        print("✅ Папка 'Шаблоны' скопирована")
    else:
        # Создаем пустую папку шаблонов
        os.makedirs(os.path.join(portable_dir, 'Шаблоны'), exist_ok=True)
        create_sample_template(os.path.join(portable_dir, 'Шаблоны'))

    # Создаем bat файлы
    create_bat_files(portable_dir)
    create_readme(portable_dir)

    print(f"📦 ПОРТАТИВНОЕ РЕШЕНИЕ СОЗДАНО: {portable_dir}")


def create_bat_files(portable_dir):
    """Создать bat файлы для портативной версии"""

    # install_dependencies.bat
    install_bat = '''@echo off
chcp 65001
title DocumentFiller - Установка зависимостей
echo ========================================
echo    Установка зависимостей DocumentFiller
echo ========================================
echo.

echo 📦 Проверка Python...
python --version >nul 2>&1
if errorlevel 1 (
    echo ❌ Python не установлен или не добавлен в PATH
    echo 📥 Скачайте Python с https://python.org
    pause
    exit /b 1
)

echo ✅ Python обнаружен
echo.

echo 🔄 Обновление pip...
python -m pip install --upgrade pip
if errorlevel 1 (
    echo ⚠️ Не удалось обновить pip, продолжаем...
)

echo.
echo 📥 Установка необходимых библиотек...
pip install -r requirements.txt
if errorlevel 1 (
    echo ❌ Ошибка установки зависимостей
    pause
    exit /b 1
)

echo.
echo ✅ Все зависимости успешно установлены!
echo.
echo 🚀 Теперь вы можете запустить программу через start.bat
echo.
pause
'''

    # start.bat
    start_bat = '''@echo off
chcp 65001
title DocumentFiller
echo ============================
echo      DocumentFiller
echo ============================
echo.

echo 🔍 Проверка зависимостей...
python -c "import PyQt5, openpyxl, docxtpl, requests" >nul 2>&1
if errorlevel 1 (
    echo ❌ Зависимости не установлены
    echo 📥 Запустите install_dependencies.bat
    pause
    exit /b 1
)

echo ✅ Зависимости проверены
echo 🚀 Запуск программы...
echo.

python main.py

echo.
echo Программа завершена
pause
'''

    # requirements.txt
    requirements = '''PyQt5>=5.15
openpyxl>=3.0
python-docx>=0.8
docxtpl>=0.16
jinja2>=3.0
lxml>=4.6
requests>=2.25
'''

    with open(os.path.join(portable_dir, 'install_dependencies.bat'), 'w', encoding='utf-8') as f:
        f.write(install_bat)

    with open(os.path.join(portable_dir, 'start.bat'), 'w', encoding='utf-8') as f:
        f.write(start_bat)

    with open(os.path.join(portable_dir, 'requirements.txt'), 'w', encoding='utf-8') as f:
        f.write(requirements)

    print("✅ Созданы bat файлы и requirements.txt")


def create_readme(dist_dir):
    """Создать README файл"""
    readme_content = '''DocumentFiller - Программа для заполнения документов

🎯 НАЗНАЧЕНИЕ:
Автоматическое заполнение шаблонов документов Word на основе данных из Excel

🚀 ЗАПУСК:

Для собранной версии (EXE):
- Запустите DocumentFiller.exe

Для портативной версии:
1. Убедитесь, что установлен Python 3.8+
2. Запустите install_dependencies.bat (только при первом запуске)
3. Запустите start.bat

📁 СТРУКТУРА ПАПОК:
- Шаблоны/ - папка с шаблонами документов (.docx)
- документы/ - папка для сохранения заполненных документов (создается автоматически)

⚙️ ТРЕБОВАНИЯ:
- Windows 7/8/10/11
- Для портативной версии: Python 3.8+

📞 ПОДДЕРЖКА:
Разработчик: Строчков Сергей Константинович
Телефон: 8(920)791-30-43

⚠️  ПРИМЕЧАНИЕ:
Перед использованием поместите ваши шаблоны документов в папку "Шаблоны"
'''

    readme_path = os.path.join(dist_dir, 'README.txt')
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(readme_content)

    print(f"📖 Создан README.txt")


if __name__ == "__main__":
    print("=" * 60)
    print("🔧 DOCUMENTFILLER - СБОРКА С ИСПРАВЛЕНИЕМ МОДУЛЯ EMAIL")
    print("=" * 60)

    build_complete()

    print("\n" + "=" * 60)
    print("🏁 Процесс сборки завершен!")
    print("=" * 60)